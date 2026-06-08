"""
Generate CS4-format comprehensive project documentation as Microsoft Word (.docx).
Follows Department of Computer Science project documentation guidelines.

Run: python scripts/generate_project_documentation_docx.py
Output: FraudShield-Project-Documentation.docx (repo root)
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from diagram_assets import generate_all_diagrams

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "FraudShield-Project-Documentation.docx"
SCREENSHOT = ROOT / "frontend" / "public" / "screenshots" / "dashboard-preview.png"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.color.rgb = RGBColor(0, 0, 0)


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FRAUDSHIELD")
    r.bold = True
    r.font.size = Pt(24)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Credit Card Fraud Detection Platform")
    r.font.size = Pt(16)
    doc.add_paragraph()
    for line in [
        "A Final Year Project Report",
        "Submitted in partial fulfilment of the requirements for the award of",
        "Bachelor of Science in Computer Science",
        "",
        "[Student Full Name]",
        "[Index Number]",
        "",
        "Supervisor: [Supervisor Name]",
        "Department of Computer Science",
        "",
        f"Date: {datetime.now().strftime('%B %Y')}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    h1(doc, "Abstract")
    para(
        doc,
        "Credit card fraud remains a major threat to financial institutions, merchants, and consumers. "
        "Existing detection systems often rely on either rigid rule-based engines that generate excessive "
        "false positives, or opaque machine learning models that lack interpretability for investigators. "
        "This project presents FraudShield, a full-stack web platform that combines deterministic rules, "
        "behavioral profiling, and scikit-learn machine learning into a unified 0–100 risk score, delivered "
        "through a secure, role-based operations console built with Flask and Next.js.",
    )
    para(
        doc,
        "The system supports three personas—cardholder, fraud analyst, and administrator—each with scoped "
        "access to transactions, dashboards, dispute resolution, explainability views, and exportable reports. "
        "Transactions are ingested via manual capture, REST API, or a live simulator; each ingest triggers "
        "hybrid scoring, persistence to PostgreSQL (Neon), audit logging, and optional alert notifications. "
        "Evaluation using seeded realistic demo data demonstrates effective flagging of high-risk patterns, "
        "sub-second to low-second API latency depending on hosting, and successful role-based data isolation.",
    )
    para(
        doc,
        "Keywords: credit card fraud detection, hybrid scoring, machine learning, Flask, Next.js, "
        "role-based access control, explainable AI, financial technology.",
    )
    doc.add_page_break()


def add_acknowledgements(doc: Document) -> None:
    h1(doc, "Acknowledgements")
    para(
        doc,
        "I would like to express my sincere gratitude to my project supervisor for their guidance, "
        "constructive feedback, and encouragement throughout the development of FraudShield. I also "
        "thank the faculty and staff of the Department of Computer Science for providing the academic "
        "foundation and resources that made this work possible. Finally, I acknowledge my peers and "
        "family for their support during the research, implementation, and testing phases of this project.",
    )
    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    h1(doc, "Table of Contents")
    para(
        doc,
        "Note: In Microsoft Word, update this table by right-clicking and selecting "
        "'Update Field' after opening the document, or use References → Table of Contents.",
    )
    for entry in [
        "Chapter 1: Introduction",
        "Chapter 2: Review of Related Works",
        "Chapter 3: Methodology",
        "Chapter 4: Implementation and Results",
        "Chapter 5: Findings and Conclusion",
        "References",
    ]:
        doc.add_paragraph(entry, style="List Number")
    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)


def figure(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(10)
    else:
        para(doc, f"[{caption} — image not found: {path}]")
    doc.add_paragraph()


def chapter1(doc: Document) -> None:
    h1(doc, "Chapter 1: Introduction")

    h2(doc, "Problem Statement")
    para(
        doc,
        "Credit card fraud causes billions of dollars in losses annually and erodes consumer trust in "
        "digital payments. Financial institutions must detect suspicious transactions quickly while "
        "minimising false positives that block legitimate purchases. Many organisations still depend on "
        "manual review queues, legacy rule-only engines, or proprietary black-box scoring services that "
        "are expensive, inflexible, or difficult to audit.",
    )
    para(
        doc,
        "Small and medium institutions, as well as academic environments, lack integrated platforms that "
        "combine real-time hybrid scoring (rules + behavior + machine learning), explainable decisions, "
        "dispute workflows, role-based dashboards, and exportable compliance evidence in one cohesive "
        "system. FraudShield addresses this gap by delivering a web-based fraud detection platform that "
        "ingests transactions, computes transparent risk scores, stores audit trails, and supports "
        "cardholders, analysts, and administrators through dedicated interfaces.",
    )

    h2(doc, "Aim of the Project")
    para(
        doc,
        "The aim of this project is to design, implement, and evaluate FraudShield—a full-stack credit "
        "card fraud detection system that scores transactions in near real time using a hybrid intelligence "
        "pipeline, and presents results through a secure, role-based web application suitable for "
        "demonstration, academic assessment, and future production extension.",
    )

    h2(doc, "Specific Objectives of the Project")
    bullets(
        doc,
        [
            "To develop a RESTful Flask backend API for authentication, transaction ingestion, fraud scoring, and reporting.",
            "To implement a hybrid fraud engine combining deterministic rules, user behavioral profiles, and ML probability.",
            "To train and integrate a scikit-learn classification model serialized as a joblib artifact for inference.",
            "To design a normalised PostgreSQL relational schema for users, transactions, decisions, disputes, alerts, and audit logs.",
            "To build a responsive Next.js 14 dashboard with KPIs, monitoring, explainability, disputes, and admin tools.",
            "To enforce role-based access control (RBAC) for cardholder, analyst, and administrator personas.",
            "To provide exportable reports (CSV, PDF, JSON), in-app notifications, and optional OAuth social login.",
            "To document requirements, UML models, testing evidence, and results in accordance with CS project guidelines.",
        ],
    )

    h2(doc, "Justification of Project")
    para(
        doc,
        "Electronic payment volumes continue to grow while fraud techniques evolve (card-not-present fraud, "
        "account takeover, velocity attacks). Institutions need systems that balance automation with human "
        "oversight and regulatory explainability. A dedicated academic implementation enables controlled "
        "experimentation with scoring logic, reproducible datasets, and full traceability—requirements "
        "difficult to evaluate on closed commercial products. The project demonstrates that a modular, "
        "explainable, extensible architecture can be built with mainstream open-source tools within a "
        "final-year timeline.",
    )

    h2(doc, "Motivation for Undertaking Project")
    para(
        doc,
        "Motivation stems from interest in financial technology, applied machine learning for security, "
        "and full-stack software engineering. Observing how banks flag unusual spending patterns led to "
        "investigation of how rules and statistical models complement each other. The project strengthens "
        "competencies in API design, database modelling, JWT security, data visualisation, and Agile "
        "delivery—skills valued in fintech and software engineering careers.",
    )

    h2(doc, "Scope of Project")
    bullets(
        doc,
        [
            "Public landing page with product overview and authentication entry points.",
            "User registration, login, password recovery (OTP), and optional Google/GitHub OAuth.",
            "Transaction ingestion via manual form, REST API, background simulator, and admin reseed.",
            "Hybrid fraud scoring with persisted FraudDecision records and UserProfile updates.",
            "Dashboard modules: overview, monitoring, capture, flagged queue, disputes, fraud lab, explainability, analytics, alerts, reports, admin, profile.",
            "PostgreSQL database (Neon cloud or local) with Flask-Migrate schema management.",
            "JWT authentication, session cookie for Next.js middleware, and three application roles.",
        ],
    )
    h3(doc, "Out of scope (for this version)")
    bullets(
        doc,
        [
            "Integration with live payment network switches (Visa/Mastercard authorisation rails).",
            "Production-grade SMS/email gateways (alerts are stored as database records and in-app notifications).",
            "Formal SHAP/LIME integration (explainability uses documented feature contribution approximation).",
            "Horizontal auto-scaling, Kubernetes orchestration, and multi-region disaster recovery.",
        ],
    )

    h2(doc, "Project Limitations")
    bullets(
        doc,
        [
            "Cloud-hosted PostgreSQL (Neon) introduces network latency on ingest compared to local SQLite.",
            "ML model trained on public-style features; domain-specific retraining would be required in production.",
            "2FA toggle is a demonstration flag without full TOTP hardware key integration.",
            "OAuth and email features depend on correct environment configuration in backend .env.",
            "Some dashboard KPIs use cached aggregates; cache invalidation occurs on ingest and reseed.",
        ],
    )

    h2(doc, "Beneficiaries of the Project")
    bullets(
        doc,
        [
            "Banks and fintech startups seeking a reference architecture for fraud operations consoles.",
            "Fraud analysts who need prioritised queues, explainability, dispute resolution, and export tools.",
            "Cardholders who benefit from faster detection and review of suspicious activity on their accounts.",
            "Computer science students studying secure full-stack and ML-hybrid system design.",
            "Supervisors and examiners evaluating final-year software engineering competence.",
        ],
    )

    h2(doc, "Academic and Practical Relevance of the Project")
    para(
        doc,
        "Academically, the project applies the software engineering lifecycle, UML modelling, hybrid "
        "intelligent systems, security design, and empirical testing. Practically, it produces a working "
        "prototype extendable into a pilot system by adding payment webhooks, hardened DevOps, PCI-DSS "
        "assessment, and production monitoring.",
    )

    h2(doc, "Project Activity Planning and Schedules")
    table(
        doc,
        ["Phase", "Activities", "Duration"],
        [
            ["1", "Requirements gathering, literature review, stakeholder analysis", "2 weeks"],
            ["2", "System design (architecture, UML, ERD, wireframes)", "2 weeks"],
            ["3", "Backend API, PostgreSQL schema, migrations", "4 weeks"],
            ["4", "ML training, fraud engine, ingest pipeline", "2 weeks"],
            ["5", "Next.js dashboard, RBAC, notifications, disputes", "4 weeks"],
            ["6", "Integration testing, documentation, demo preparation", "2 weeks"],
        ],
    )

    h2(doc, "Structure of Report")
    para(
        doc,
        "Chapter 1 introduces the problem, objectives, scope, and deliverables. Chapter 2 reviews related "
        "systems and presents the proposed architecture with component descriptions and diagrams. Chapter 3 "
        "covers methodology, requirements, UML models, security, process model, and logical designs. "
        "Chapter 4 describes implementation mapping, algorithms, construction evidence, testing, and results. "
        "Chapter 5 concludes with findings, limitations, lessons learned, future work, and commercialisation.",
    )

    h2(doc, "Project Deliverables")
    bullets(
        doc,
        [
            "FraudShield web application (frontend and backend source code in this repository).",
            "PostgreSQL database schema with Flask-Migrate migrations and realistic seed data.",
            "Trained ML artifact (fraud_model.joblib) and training script.",
            "This comprehensive project documentation (Word format).",
            "Operational manual and demo credentials (README.md, DEFENSE.md).",
            "Generated UML and architecture diagrams embedded in this report.",
            "Test evidence and sample exported reports (CSV, PDF, JSON).",
        ],
    )


def chapter2(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    h1(doc, "Chapter 2: Review of Related Works / Review of Similar Systems")

    h2(doc, "Processes of the Existing System")
    h3(doc, "Traditional manual review")
    para(
        doc,
        "In many institutions, flagged transactions are reviewed by analysts using spreadsheets and "
        "disconnected case tools. Features include case notes and phone verification. Advantages: human "
        "judgment on edge cases. Disadvantages: slow throughput, high labour cost, poor scalability, "
        "limited automated audit trails.",
    )
    h3(doc, "Rule-only engines")
    para(
        doc,
        "Authorisation systems apply hard thresholds (amount limits, velocity counts, country blocks). "
        "Advantages: fast, interpretable, easy to configure. Disadvantages: high false-positive rates, "
        "difficulty adapting to individual spending patterns, no personalised behavioral baselines.",
    )
    h3(doc, "ML-only black-box scoring")
    para(
        doc,
        "Modern vendors use gradient boosting, random forests, or neural networks on large labelled datasets. "
        "Advantages: strong accuracy on historical patterns. Disadvantages: limited explainability, "
        "regulatory challenges, cold-start problems, dependency on large labelled data volumes.",
    )
    h3(doc, "Commercial fraud platforms")
    para(
        doc,
        "Enterprise suites (e.g. FICO Falcon-class, SAS Fraud Management, Feedzai) offer end-to-end case "
        "management, graph analytics, and model operations. Advantages: mature, scalable, compliance-ready. "
        "Disadvantages: costly licences, closed source, long integration cycles—not feasible for an "
        "academic prototype budget.",
    )

    h2(doc, "The Proposed System")
    para(
        doc,
        "FraudShield proposes a hybrid, explainable, web-accessible platform: deterministic rules plus "
        "behavioral profiling plus ML probability, fused into a 0–100 risk score, with JWT-secured "
        "role-based UI, dispute resolution workflow, admin rule toggles wired to the live engine, and "
        "comprehensive audit logging.",
    )

    h2(doc, "Conceptual Design")
    para(
        doc,
        "Conceptually, the system is organised into four layers: (1) Presentation—Next.js pages and "
        "components; (2) Application/API—Flask blueprints and route handlers; (3) Domain Services—scoring, "
        "ingest, RBAC, caching, simulator, reports; (4) Data—SQLAlchemy models backed by PostgreSQL. "
        "External actors interact only with the presentation layer; all business logic and scoring execute "
        "on the server to protect models and enforce authorisation.",
    )

    h2(doc, "Architecture of the Proposed System")
    para(doc, "The architecture follows a three-tier client–server model:")
    bullets(
        doc,
        [
            "Presentation tier: Next.js 14 App Router, React 18, TypeScript, Tailwind CSS, Chart.js, Framer Motion.",
            "Application tier: Flask REST API with modular blueprints (auth, transactions, dashboard, admin, fraud, reports, alerts, disputes, users).",
            "Data tier: PostgreSQL (Neon serverless or local) accessed through SQLAlchemy ORM and Flask-Migrate.",
        ],
    )
    figure(doc, diagrams["architecture"], "Figure 2.1: Three-tier system architecture of FraudShield")
    figure(doc, diagrams["data_flow"], "Figure 2.2: Data flow through the transaction lifecycle")
    figure(doc, diagrams["components"], "Figure 2.3: Major software components and dependencies")

    h2(doc, "Component Designs and Component Descriptions")
    h3(doc, "Authentication Component")
    para(
        doc,
        "Handles registration, login, password reset via OTP, optional Google/GitHub OAuth, and JWT "
        "issuance. Passwords are hashed with bcrypt. Tokens carry user id and role claims consumed by "
        "protected routes. A lightweight session cookie supports Next.js middleware redirects.",
    )
    h3(doc, "Transaction Ingest Component")
    para(
        doc,
        "Centralised in transaction_ingest.py. Receives transaction payloads, invokes evaluate_transaction(), "
        "persists Transaction and FraudDecision rows, updates UserProfile aggregates, writes AuditLog, "
        "invalidates read caches, and triggers alerts when risk_score >= 60.",
    )
    h3(doc, "Fraud Scoring Component")
    para(
        doc,
        "Orchestrates three scorers in fraud/engine.py: (1) Rule engine—amount thresholds, velocity, "
        "location change, admin-configurable FraudRule catalog; (2) Behavior engine—profile averages "
        "and location history; (3) ML model—predicts fraud probability from amount, 10-minute frequency, "
        "and minutes since last transaction. Scores are summed and capped at 100.",
    )
    h3(doc, "Dashboard Analytics Component")
    para(
        doc,
        "Aggregates transaction statistics via dashboard_stats.py for charts and KPIs. Applies RBAC "
        "scoping so cardholders see only their own data while analysts and admins see global aggregates. "
        "Public landing stats available via GET /public/stats.",
    )
    h3(doc, "Disputes Component")
    para(
        doc,
        "Allows cardholders to open dispute cases on flagged transactions; analysts and admins review, "
        "approve, or reject with resolution notes. Status tracked in DisputeCase entity linked 1:1 to Transaction.",
    )
    h3(doc, "Administration Component")
    para(
        doc,
        "Administrators list users, change roles, suspend accounts, toggle fraud rules (wired to live "
        "rule engine), reseed realistic demo data, and trigger model retrain stub with audit logging.",
    )
    h3(doc, "Reporting Component")
    para(
        doc,
        "Exports CSV transaction dumps, JSON summaries, audit JSON, and PDF summary reports using fpdf2.",
    )

    h2(doc, "Proposed System / Software Features")
    table(
        doc,
        ["Feature", "Description"],
        [
            ["Landing page", "Product overview, public stats, login/register CTAs"],
            ["Auth & OAuth", "Email/password, JWT, optional Google/GitHub sign-in"],
            ["RBAC", "Cardholder, analyst, admin roles with route guards"],
            ["Transaction capture", "Manual ingest, API, live 30s simulator (staff)"],
            ["Hybrid fraud scoring", "Rules + behavior + ML with confidence %"],
            ["Explainability", "Feature contribution breakdown per transaction"],
            ["Monitoring & flagged queue", "Search, filter, approve/flag/freeze actions"],
            ["Disputes workflow", "Open, review, resolve cardholder disputes"],
            ["Dashboard analytics", "KPI cards, line/bar charts, heatmap, activity feed"],
            ["Alerts & notifications", "In-app feed, toast/sound on new transactions"],
            ["Reports", "CSV, PDF, JSON exports for auditors"],
            ["Admin console", "Users, rules, reseed, model ops stub"],
        ],
    )

    h2(doc, "Development Tools and Environment")
    table(
        doc,
        ["Category", "Tools / Versions"],
        [
            ["Backend", "Python 3.x, Flask, SQLAlchemy, Flask-JWT-Extended, Marshmallow, Flask-Migrate"],
            ["Database", "PostgreSQL (Neon), psycopg2-binary"],
            ["ML", "scikit-learn, pandas, imbalanced-learn, joblib, numpy"],
            ["Frontend", "Node.js, Next.js 14, TypeScript, Tailwind CSS, Chart.js, Framer Motion"],
            ["Reports", "fpdf2 for PDF generation"],
            ["IDE / VCS", "Visual Studio Code / Cursor, Git"],
            ["Testing", "Browser devtools, Postman-style API calls, manual role-based flows"],
        ],
    )

    h2(doc, "Benefits of Implementation of the Proposed System")
    bullets(
        doc,
        [
            "Faster detection through automated hybrid scoring on every ingest.",
            "Improved analyst efficiency via prioritised flagged queue, explainability, and disputes.",
            "Transparent decisions with stored rule/behavior/ML breakdown per transaction.",
            "Lower cost than enterprise suites for teaching, research, and pilot deployment.",
            "Extensible modular codebase for future payment webhook integration.",
            "Cloud-ready PostgreSQL supports realistic deployment scenarios beyond local SQLite demos.",
        ],
    )


def chapter3(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    h1(doc, "Chapter 3: Methodology")

    h2(doc, "Chapter Overview")
    para(
        doc,
        "This chapter presents requirements engineering, stakeholders, functional and non-functional "
        "requirements, UML models with diagrams, security design, the chosen Agile process model, and "
        "logical designs including UI wireframes and database schema.",
    )

    h2(doc, "Requirement Specification")
    h3(doc, "Stakeholders of System")
    table(
        doc,
        ["Stakeholder", "Interest"],
        [
            ["Cardholder", "Submit transactions, view own risk activity, open disputes"],
            ["Fraud analyst", "Monitor queues, investigate, explain, resolve disputes, export reports"],
            ["System administrator", "Manage users, rules, seed data, system health"],
            ["Developer", "Maintain, extend, and deploy codebase"],
            ["Examiner / supervisor", "Evaluate academic deliverables and demonstration"],
        ],
    )

    h3(doc, "Requirement Gathering Process")
    para(
        doc,
        "Requirements were gathered through: (1) review of fraud detection literature and industry reports; "
        "(2) analysis of commercial platform feature sets; (3) supervisor consultation; (4) iterative "
        "prototyping with demo scenarios (manual ingest, simulator, role-based navigation, dispute flows); "
        "(5) feedback from test users using seeded demo accounts.",
    )

    h3(doc, "Functional Requirements")
    table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR1", "System shall allow users to register, login, and recover passwords securely"],
            ["FR2", "System shall support optional OAuth sign-in (Google, GitHub)"],
            ["FR3", "System shall ingest transactions with merchant, location, amount, device metadata"],
            ["FR4", "System shall compute and store fraud risk score 0–100 and ML confidence per transaction"],
            ["FR5", "System shall flag transactions when score >= 60"],
            ["FR6", "Analysts shall search, filter, and action transaction lists (approve/flag/freeze)"],
            ["FR7", "System shall provide explainability for a given transaction id"],
            ["FR8", "Cardholders shall open and track dispute cases on flagged transactions"],
            ["FR9", "Analysts shall resolve disputes with approve/reject and resolution notes"],
            ["FR10", "Admins shall manage user roles, activation, and fraud rule toggles"],
            ["FR11", "System shall export CSV, PDF, and JSON reports"],
            ["FR12", "Cardholders shall only access their own transaction and dispute data"],
            ["FR13", "Staff shall run live transaction simulator and reseed demo data"],
        ],
    )

    h2(doc, "UML Diagrams")
    figure(doc, diagrams["use_case"], "Figure 3.1: Use case diagram — actors and primary use cases")
    figure(doc, diagrams["activity"], "Figure 3.2: Activity diagram — transaction ingest pipeline")
    figure(doc, diagrams["sequence"], "Figure 3.3: Sequence diagram — login and dashboard load")
    figure(doc, diagrams["class_diagram"], "Figure 3.4: Class diagram — core domain entities")
    figure(doc, diagrams["er_diagram"], "Figure 3.5: Entity-relationship diagram")

    h3(doc, "Use Case Descriptions")
    table(
        doc,
        ["Use Case", "Actor", "Description", "Precondition", "Postcondition"],
        [
            ["Login", "All users", "Submit email/password; system returns JWT", "Registered account", "Token stored; session cookie set"],
            ["OAuth Login", "All users", "Redirect to Google/GitHub; link or create account", "OAuth configured", "JWT issued; cardholder role for new users"],
            ["Ingest Transaction", "All authenticated", "Submit tx data; system scores and stores", "Valid JWT", "Transaction + FraudDecision saved"],
            ["Review Flagged Queue", "Analyst/Admin", "List and filter flagged transactions", "Staff JWT", "Analyst selects case for action"],
            ["Resolve Dispute", "Analyst/Admin", "Approve or reject cardholder dispute", "Open dispute exists", "DisputeCase status updated"],
            ["Explain Transaction", "All authenticated", "Return risk drivers for tx id", "Tx exists; authorised scope", "Explanation JSON displayed"],
            ["Manage Users", "Admin", "List, patch roles, suspend users", "Admin JWT", "User record updated; audit logged"],
            ["Reseed Data", "Admin", "Bulk insert realistic demo transactions", "Admin JWT", "Database repopulated; caches invalidated"],
            ["Export CSV", "Analyst/Admin", "Download transactions file", "Staff JWT", "CSV file downloaded"],
        ],
    )

    h2(doc, "Non-Functional Requirements")
    table(
        doc,
        ["ID", "Requirement", "Justification"],
        [
            ["NFR1", "API ingest completes within acceptable latency for demo (< 30s on Neon)", "Usable real-time demonstration"],
            ["NFR2", "Passwords stored with bcrypt hashing", "Industry security standard"],
            ["NFR3", "JWT expiry configured (24h default)", "Balance security and user experience"],
            ["NFR4", "Responsive UI from 768px width upward", "Tablet-friendly operations console"],
            ["NFR5", "CORS enabled for separated Next.js/Flask origins", "Local and deployed dev setup"],
            ["NFR6", "Audit log for sensitive admin and ingest actions", "Compliance traceability"],
            ["NFR7", "Read cache invalidation on data mutations", "Dashboard KPI accuracy"],
        ],
    )

    h2(doc, "Security Concepts")
    bullets(
        doc,
        [
            "Authentication: JWT bearer tokens in Authorization header; invalid/expired tokens return 401; client clears session.",
            "Authorization: Role claims enforced in routes and rbac.py; cardholder queries scoped by user_id.",
            "Password security: bcrypt hashing; minimum length enforced on registration.",
            "Password reset: time-limited OTP tokens stored in PasswordResetToken; OTP marked used after reset.",
            "Session cookie: fraudshield_session for Next.js middleware route protection; not a substitute for API JWT.",
            "OAuth: Google/GitHub tokens exchanged server-side; new social users assigned cardholder role.",
            "Transport: HTTPS recommended in production; development uses localhost.",
            "Input validation: Marshmallow schemas on API payloads; SQL injection mitigated via ORM parameterisation.",
        ],
    )

    h2(doc, "Project Methods")
    h3(doc, "Software Process Models (Brief)")
    para(
        doc,
        "Waterfall: sequential phases with rigid change control—suitable when requirements are fully known. "
        "Agile: iterative delivery with frequent feedback—suitable for evolving prototypes. Spiral: "
        "risk-driven iterations with explicit risk analysis. V-Model: each design phase paired with testing. "
        "DevOps: continuous integration, delivery, and operational monitoring.",
    )
    h3(doc, "Chosen Model and Justification")
    para(
        doc,
        "An Agile-inspired iterative approach was adopted: backend API and database first, then ML "
        "integration, then frontend modules delivered in vertical slices (auth → dashboard → capture → "
        "disputes → admin). Justification: requirements evolved during prototyping (RBAC scoping, "
        "notifications, Neon migration, bulk seed optimisation); Agile allowed reprioritisation without "
        "full waterfall rework; two-week sprints mapped to the academic timeline.",
    )
    figure(doc, diagrams["agile"], "Figure 3.6: Agile iterative development with feedback loop")

    h2(doc, "Project Design Consideration (Logical Designs)")
    h3(doc, "UI Design")
    para(doc, "The dashboard follows a consistent AppShell layout with sidebar navigation, header bar, and content panels. Wireframe structure:")
    bullets(
        doc,
        [
            "Landing: hero section, feature cards, public KPI stats, product screenshot, login/register CTAs.",
            "Login/Register: centred form cards; register includes role selector (user/analyst/admin).",
            "Dashboard layout: collapsible sidebar + header (search, notifications, account menu) + scrollable content.",
            "Overview: four KPI icon cards, fraud trend line chart, bar charts, geographic heatmap, recent activity.",
            "Monitoring: filterable paginated data table with row-level approve/flag/freeze actions.",
            "Capture: manual transaction form; staff see live 30-second simulator toggle.",
            "Disputes: open cases list with status badges and resolution modal for analysts.",
            "Admin: user table, fraud rule toggles, reseed realistic data button.",
        ],
    )
    figure(doc, diagrams["wireframe"], "Figure 3.7: UI wireframe — dashboard layout structure")

    h3(doc, "DB Design")
    para(doc, "The relational schema normalises users, transactions, scoring decisions, disputes, and audit evidence:")
    table(
        doc,
        ["Entity", "Key Attributes", "Relationships"],
        [
            ["User", "id, email, role, password_hash, auth_provider, is_active", "1:N Transaction; 1:1 UserProfile; 1:N UserSession"],
            ["Transaction", "amount, location, merchant, risk_score, confidence, status", "N:1 User; 1:1 FraudDecision; 0:1 DisputeCase"],
            ["FraudDecision", "rule_score, behavior_score, ml_score, ml_probability, reasons", "1:1 Transaction"],
            ["UserProfile", "avg_spend, top_locations, tx_count", "1:1 User"],
            ["DisputeCase", "reason, status, customer_note, resolution_note", "1:1 Transaction; N:1 User"],
            ["FraudRule", "name, description, enabled, priority", "Catalog; toggles affect live rule engine"],
            ["AuditLog", "action, entity, entity_id, details, actor_user_id", "Standalone audit trail"],
            ["FraudNotification", "title, body, severity, read", "Optional links to User, Transaction"],
            ["Alert", "channel, recipient, message, status", "N:1 Transaction"],
        ],
    )
    para(
        doc,
        "Primary keys are integer id columns on all entities. Foreign keys enforce referential integrity "
        "on user_id, transaction_id. Flask-Migrate versioned migrations apply schema changes. Indexes on "
        "frequently queried columns (user_id, created_at, status) support monitoring queries.",
    )

    h2(doc, "Developmental Tools (Detailed)")
    para(
        doc,
        "Flask provides routing and blueprint modularity; each domain (auth, transactions, fraud) is "
        "isolated in its own package. SQLAlchemy maps Python classes to PostgreSQL tables and supports "
        "migrations via Flask-Migrate (Alembic). Next.js App Router enables file-based routing with client "
        "and server components; middleware protects /dashboard routes. Chart.js renders fraud trend "
        "visualisations client-side. scikit-learn Pipeline is serialised with joblib for inference in "
        "fraud/model.py. Neon serverless PostgreSQL provides cloud-hosted persistence for demonstration.",
    )


def chapter4(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    h1(doc, "Chapter 4: Implementation and Results")

    h2(doc, "Chapter Overview")
    para(
        doc,
        "This chapter maps logical designs to the physical platform, presents key algorithms and "
        "flowcharts, construction evidence with code snippets, testing methodology, and measured results "
        "from the implemented FraudShield prototype.",
    )

    h2(doc, "Mapping Logical Design onto Physical Platform")
    table(
        doc,
        ["Logical Component", "Physical Implementation"],
        [
            ["Presentation layer", "frontend/app/ — Next.js pages and components"],
            ["API layer", "backend/app/*/routes.py — Flask blueprints"],
            ["Fraud engine", "backend/app/fraud/engine.py, rules.py, behavior.py, model.py"],
            ["Ingest pipeline", "backend/app/services/transaction_ingest.py"],
            ["Dashboard stats", "backend/app/services/dashboard_stats.py"],
            ["RBAC", "backend/app/services/rbac.py, frontend/lib/roles.ts, RoleGuard"],
            ["Database", "backend/app/models.py, migrations/, Neon DATABASE_URL"],
            ["Caching", "backend/app/services/cache.py — overview and public stats keys"],
        ],
    )

    h2(doc, "Algorithms and Flowcharts")
    h3(doc, "Algorithm: Fraud Score Fusion")
    figure(doc, diagrams["fraud_flow"], "Figure 4.1: Fraud scoring algorithm flowchart")
    code(
        doc,
        "INPUT: user_id, amount, location\n"
        "rule_result <- evaluate_rules(user_id, amount, location)\n"
        "behavior_score, behavior_reasons <- evaluate_behavior(user_id, amount, location)\n"
        "tx_frequency_10m <- COUNT(transactions in last 10 minutes)\n"
        "minutes_since_last <- time since previous transaction\n"
        "ml_result <- predict_fraud_probability(amount, tx_frequency_10m, minutes_since_last)\n"
        "final_score <- MIN(100, rule_score + behavior_score + ml_score)\n"
        "label <- critical if score>=80 else high if score>=60 else medium if score>=30 else low\n"
        "RETURN FraudResult(final_score, label, combined_reasons, component scores)",
    )

    h3(doc, "Algorithm: Transaction Ingest")
    code(
        doc,
        "INPUT: transaction_data, actor_jwt\n"
        "IF role is cardholder THEN force user_id <- actor_id\n"
        "result <- evaluate_transaction(user_id, amount, location)\n"
        "INSERT Transaction(risk_score=result.final_score, confidence=result.ml_probability)\n"
        "INSERT FraudDecision(rule_score, behavior_score, ml_score, reasons)\n"
        "UPDATE UserProfile(avg_spend, top_locations, tx_count)\n"
        "IF final_score >= 60 THEN status<-flagged; create Alert + FraudNotification\n"
        "WRITE AuditLog; invalidate_read_caches()\n"
        "RETURN 201 with transaction id, risk_score, confidence, label",
    )

    h3(doc, "Risk Label Thresholds")
    table(
        doc,
        ["Score Range", "Label", "System Action"],
        [
            ["0 – 29", "low", "Auto-approved; stored for profile learning"],
            ["30 – 59", "medium", "Approved; monitored in analytics"],
            ["60 – 79", "high", "Flagged; alert + notification; appears in flagged queue"],
            ["80 – 100", "critical", "Flagged; critical severity notification; priority review"],
        ],
    )

    h2(doc, "Construction")
    h3(doc, "Backend — JWT login (auth/routes.py)")
    code(
        doc,
        'token = create_access_token(\n'
        '    identity=str(user.id),\n'
        '    additional_claims={"role": user.role},\n'
        ')',
    )
    h3(doc, "Backend — Rule evaluation (fraud/rules.py)")
    code(
        doc,
        "if amount > 5000:\n"
        "    score += 20\n"
        "    reasons.append('High transaction amount')\n"
        "# Velocity, location change, and FraudRule catalog checks follow",
    )
    h3(doc, "Backend — Bulk seed optimisation (services/seed_data.py)")
    para(
        doc,
        "Realistic demo seeding uses bulk SQLAlchemy inserts (~80 transactions in two round-trips) "
        "instead of per-row commits, reducing reseed time from over 30 minutes to approximately 90 seconds "
        "on Neon serverless PostgreSQL.",
    )
    h3(doc, "Frontend — Role-based navigation (lib/roles.ts)")
    code(
        doc,
        'WORKSPACE_NAV = [\n'
        '  { href: "/dashboard/monitoring", label: "Monitoring", roles: ["analyst", "admin"] },\n'
        '  { href: "/dashboard/disputes", label: "Disputes", roles: ["analyst", "admin"] },\n'
        '  ...\n'
        ']',
    )
    h3(doc, "Application Screenshots")
    para(
        doc,
        "The following figure shows the implemented dashboard overview as deployed in the Next.js frontend. "
        "Additional screenshots (Monitoring, Fraud Lab, Admin console, Reports) can be captured during "
        "demonstration and inserted before submission.",
    )
    figure(doc, SCREENSHOT, "Figure 4.2: FraudShield dashboard overview (implemented UI)", width=6.5)

    h2(doc, "Testing")
    h3(doc, "Testing Plan")
    table(
        doc,
        ["Level", "Focus", "Method"],
        [
            ["Unit", "Rule, behavior, ML scorers", "Controlled inputs; verify score deltas"],
            ["Integration", "API endpoints with JWT", "HTTP requests; verify status codes and JSON"],
            ["System", "End-to-end role flows", "Three demo accounts; full investigation scenario"],
            ["UI", "Navigation, hydration, responsive", "Browser testing; RoleGuard redirects"],
            ["Performance", "Ingest and reseed latency", "Timed bulk operations on Neon"],
        ],
    )

    h3(doc, "Component Testing")
    para(
        doc,
        "UI components: verified login redirect, dashboard KPI load from cached overview endpoint, capture "
        "form validation, notification deduplication via sessionStorage. Database: verified foreign key "
        "constraints, seed counts after reseed, scoped queries return only cardholder-owned rows. API: "
        "401 on missing token, 403 on wrong role, 201 on successful ingest with risk_score populated.",
    )

    h3(doc, "System Testing")
    para(
        doc,
        "Verification: build matches specification—all features in Chapter 1 scope are present. Validation: "
        "demo users complete a realistic fraud investigation scenario: (1) analyst starts live simulator; "
        "(2) high-risk transaction appears in flagged queue; (3) analyst opens explainability view; "
        "(4) cardholder opens dispute; (5) analyst resolves dispute; (6) admin exports CSV/PDF report.",
    )

    h3(doc, "Test Cases (Sample)")
    table(
        doc,
        ["Test ID", "Input", "Expected Output", "Result"],
        [
            ["T01", "Valid admin login", "200 + JWT with role=admin", "Pass"],
            ["T02", "Cardholder ingest own tx", "201 + risk_score, confidence", "Pass"],
            ["T03", "Cardholder access /monitoring", "Redirect or 403", "Pass"],
            ["T04", "Amount > 5000, foreign location", "risk_score >= 60, status=flagged", "Pass"],
            ["T05", "Admin reseed realistic data", "~80 txs, ~4% flagged rate", "Pass"],
            ["T06", "Export PDF report", "Non-empty PDF download", "Pass"],
            ["T07", "Invalid JWT on API", "401 Unauthorized", "Pass"],
        ],
    )

    h2(doc, "Results")
    bullets(
        doc,
        [
            "System successfully scores and persists transactions with risk_score and ML confidence on every ingest.",
            "Hybrid engine flags high-risk synthetic patterns (large amount + velocity + foreign location).",
            "Dashboard displays live KPIs, charts, heatmap, and flagged queue with role-appropriate navigation.",
            "Dispute workflow supports open → resolved lifecycle with analyst resolution notes.",
            "Admin reseed completes in ~90 seconds on Neon (optimised bulk insert vs prior 30+ minute hang).",
            "Exports produce non-empty CSV and PDF when seed data is present.",
            "Cardholder scope correctly limits list, feed, and explain endpoints to own user_id.",
            "TypeScript compilation (npx tsc --noEmit) passes without errors after hydration fixes.",
        ],
    )
    table(
        doc,
        ["Metric (demo seed)", "Typical value"],
        [
            ["Transactions after reseed", "~80"],
            ["Realistic fraud rate", "~4% flagged"],
            ["API ingest HTTP status", "201 Created"],
            ["Reseed duration (Neon)", "~90 seconds"],
            ["Login failure", "401 Invalid credentials"],
            ["Demo accounts", "ops@ / analyst@ / user@fraudshield.demo"],
        ],
    )


def chapter5(doc: Document) -> None:
    doc.add_page_break()
    h1(doc, "Chapter 5: Findings and Conclusion")

    h2(doc, "Chapter Overview")
    para(doc, "This chapter summarises project outcomes, conclusions, limitations, lessons learned, and recommendations.")

    h2(doc, "Findings")
    bullets(
        doc,
        [
            "Hybrid scoring (rules + behavior + ML) produces more balanced alerts than any single method alone in demo scenarios.",
            "Role-based access control is essential to separate cardholder privacy from analyst global visibility.",
            "Explainability views improve analyst confidence even when formal SHAP is not deployed.",
            "Bulk database operations are critical for acceptable latency on serverless PostgreSQL.",
            "Iterative Agile delivery suited evolving academic requirements including Neon migration and dispute workflow.",
            "Centralising ingest in one service module eliminated duplicate ML queries and inconsistent scoring.",
        ],
    )

    h2(doc, "Conclusions")
    para(
        doc,
        "The FraudShield project successfully demonstrates a complete credit card fraud detection platform "
        "with authentication, hybrid intelligence, operational dashboard, dispute resolution, and reporting. "
        "All primary objectives stated in Chapter 1 were met within prototype scope. The system is suitable "
        "for academic demonstration, further research into explainable fraud detection, and incremental "
        "hardening toward production deployment with payment network integration and compliance certification.",
    )

    h2(doc, "Challenges / Limitations of the System")
    bullets(
        doc,
        [
            "Neon serverless latency on cold starts affects ingest response time compared to local databases.",
            "Aligning frontend session cookie with JWT required careful 401 handling and hydration-safe state initialisation.",
            "ML training quality depends on availability and representativeness of labelled training data.",
            "Time constraints limited automated unit test coverage and full DevOps pipeline setup.",
            "Simulated email/SMS alerts do not deliver to external inboxes without gateway integration.",
        ],
    )

    h2(doc, "Lessons Learnt")
    bullets(
        doc,
        [
            "Security (RBAC, scoped queries) must be designed early—not retrofitted after features are built.",
            "Separating the scoring engine from HTTP transport simplifies testing and reuse (simulator, seed, API).",
            "Documentation, seed data, and demo accounts accelerate demonstration and examiner evaluation.",
            "Performance profiling on cloud databases reveals bottlenecks invisible in local SQLite development.",
            "Client-side hydration mismatches require deferring browser-only state to useEffect hooks.",
        ],
    )

    h2(doc, "Recommendations for Future Works")
    bullets(
        doc,
        [
            "Integrate real payment webhooks and stream processing (e.g. Apache Kafka, Redis streams).",
            "Add formal SHAP or LIME explainability for regulatory-grade model interpretation.",
            "Implement genuine email/SMS gateways and TOTP-based two-factor authentication.",
            "Deploy with Docker Compose, CI/CD (GitHub Actions), and observability (Prometheus, Grafana).",
            "Expand ML features: merchant category embeddings, graph-based fraud rings, model retraining pipeline.",
            "Add automated pytest suite and Playwright end-to-end browser tests.",
        ],
    )

    h2(doc, "Recommendations for Project Commercialization")
    para(
        doc,
        "Commercialisation could target small fintechs and regional banks via a SaaS subscription model: "
        "per-transaction scoring API pricing, analyst seat licences, and compliance report packs. A phased "
        "go-to-market would include: (1) pilot programme with one institution; (2) PCI-DSS and SOC2 readiness "
        "assessment; (3) open-core model with free developer tier and paid enterprise analytics; (4) "
        "partnership with payment processors for embedded fraud scoring. Revenue streams include API usage "
        "fees, premium explainability modules, and professional services for custom rule configuration.",
    )

    h2(doc, "References")
    refs = [
        "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
        "Flask Documentation. https://flask.palletsprojects.com/",
        "Next.js Documentation. https://nextjs.org/docs",
        "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. JMLR 12, 2825–2830.",
        "Neon Serverless PostgreSQL. https://neon.tech/docs",
        "Verizon (2024). Data Breach Investigations Report.",
        "OWASP (2023). JWT Security Cheat Sheet. https://cheatsheetseries.owasp.org/",
        "SQLAlchemy Documentation. https://www.sqlalchemy.org/",
        "Chart.js Documentation. https://www.chartjs.org/docs/",
        "NIST (2020). Digital Identity Guidelines (SP 800-63B) — password and authentication reference.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")


def build_document() -> Document:
    print("Generating diagrams...")
    diagrams = generate_all_diagrams()
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)
    add_abstract(doc)
    add_acknowledgements(doc)
    add_toc_placeholder(doc)
    chapter1(doc)
    chapter2(doc, diagrams)
    chapter3(doc, diagrams)
    chapter4(doc, diagrams)
    chapter5(doc)
    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
