"""
Generate Word document with final-year project improvement recommendations.
Run: python scripts/generate_fyp_recommendations_docx.py
Output: FraudShield-FYP-Recommendations.docx (repo root)
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "FraudShield-FYP-Recommendations.docx"


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0, 0, 0)


def title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FraudShield")
    r.bold = True
    r.font.size = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Year Project — Improvement Recommendations")
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated {datetime.now().strftime('%d %B %Y')}")
    doc.add_page_break()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table_rows(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_defaults(doc)
    title_page(doc)

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "FraudShield is already a feature-rich final-year platform: JWT authentication, role-based access "
        "control (cardholder, analyst, admin), hybrid fraud scoring (rules + behavioral analytics + machine "
        "learning), live transaction monitoring, explainability, disputes, audit logging, and exportable reports. "
        "This document prioritizes improvements that strengthen academic credibility, demo reliability, and "
        "examiner confidence — without unnecessary scope creep."
    )

    doc.add_heading("1. Fixes completed in this improvement pass", level=1)
    bullets(
        doc,
        [
            "Analyst explainability endpoint repaired — GET /fraud/explain/<id> now returns full ML and feature data for staff.",
            "Dashboard model metrics wired to ml/artifacts/metrics.json instead of hardcoded placeholder values.",
            "Admin fraud rule toggles now affect live scoring via fraud_rules_config (high_amount, velocity_window, location_mismatch, behavior_deviation).",
            "Dispute resolution workflow added — analysts can approve or reject open cases at /dashboard/disputes.",
            "ML bootstrap script added — python ml/bootstrap_model.py trains a starter model when Kaggle data is unavailable.",
            "Frontend and backend caching improved for faster navigation with remote PostgreSQL (Neon).",
        ],
    )

    doc.add_heading("2. Fix before demo (high impact)", level=1)

    doc.add_heading("2.1 Analyst explainability", level=2)
    doc.add_paragraph(
        "Examiners will ask: 'Show me why this transaction was flagged.' The explainability page and "
        "GET /fraud/explain/<transaction_id> endpoint must work for analyst and admin roles. Cardholders "
        "should continue to see simplified status only (no ML probabilities)."
    )

    doc.add_heading("2.2 Honest model metrics", level=2)
    doc.add_paragraph(
        "Present real evaluation numbers from metrics.json in the dashboard and in your written report. "
        "If PR-AUC is low on the simplified 3-feature model, explain why in the viva: class imbalance, "
        "feature reduction for real-time scoring, and fusion with rules/behavior that catch obvious fraud."
    )

    doc.add_heading("2.3 Trained ML artifact", level=2)
    bullets(
        doc,
        [
            "Run: python ml/bootstrap_model.py (synthetic bootstrap for demo)",
            "Or with Kaggle ULB dataset: python ml/train_model.py --dataset path/to/creditcard.csv",
            "Verify ml/artifacts/fraud_model.joblib exists before defense day.",
        ],
    )

    doc.add_heading("2.4 Complete dispute workflow", level=2)
    doc.add_paragraph(
        "End-to-end story: cardholder disputes → analyst reviews at Disputes page → approve/reject → "
        "customer notification → audit log entry. This demonstrates banking-style case management."
    )

    doc.add_heading("3. Strengthen the academic / ML pillar", level=1)
    table_rows(
        doc,
        ["Area", "Current state", "Recommended improvement"],
        [
            [
                "Dataset",
                "Kaggle ULB referenced; live app uses 3 engineered features",
                "In report: justify feature reduction for latency; optionally train on V1–V28 offline",
            ],
            [
                "Evaluation",
                "metrics.json with recall, precision, PR-AUC",
                "Add confusion matrix and PR curve to report appendix",
            ],
            [
                "Threshold",
                "Fixed score threshold (~60)",
                "Document threshold selection (recall vs false-positive trade-off)",
            ],
            [
                "Explainability",
                "Surrogate feature contribution bars",
                "Label as surrogate explainability; mention SHAP as future work",
            ],
            [
                "Admin rules",
                "DB toggles now affect scoring",
                "Demo: disable velocity rule and show score change on Fraud Lab",
            ],
        ],
    )

    doc.add_heading("4. Key defense sentence", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        '"We use a hybrid detector: rules catch known patterns, behavior catches profile drift, '
        "ML catches subtle anomalies — fused with explicit weights so analysts can trust and audit decisions.\""
    )
    r.italic = True

    doc.add_heading("5. Security and production (talking points)", level=1)
    bullets(
        doc,
        [
            "Rate limiting on login and transaction ingest (future work)",
            "Tighten CORS in production (currently open for development)",
            "Rotate SECRET_KEY and JWT_SECRET_KEY; never use dev defaults in production",
            "HttpOnly cookies vs localStorage JWT — document XSS trade-off",
            "Basic pytest tests for scoring engine and RBAC (5–10 tests would stand out)",
            "Optional docker-compose for Flask + Postgres + Next.js",
        ],
    )

    doc.add_heading("6. Documentation deliverables", level=1)
    bullets(
        doc,
        [
            "Architecture diagram: ingest → rules + behavior + ML → decision → alert → analyst → audit",
            "Update frontend/README.md to match actual dashboard pages",
            "Report chapter on limitations: stub email, approximate explainability, demo seeding",
            "Landing page screenshot at frontend/public/screenshots/dashboard-preview.png",
            "Evaluation appendix: metrics.json, model comparison table, explainability screenshots",
        ],
    )

    doc.add_heading("7. What NOT to add (scope creep)", level=1)
    bullets(
        doc,
        [
            "Real payment gateway integration",
            "Blockchain or buzzword features without academic justification",
            "Microservices rewrite",
            "Mobile app unless required by supervisor",
            "Real SMS/email provider (stub is acceptable if documented)",
        ],
    )

    doc.add_heading("8. Suggested two-week plan", level=1)

    doc.add_heading("Week 1 — credibility", level=2)
    bullets(
        doc,
        [
            "Fix explain endpoint and train/bootstrap model",
            "Wire real metrics; complete dispute analyst UI",
            "Architecture diagram and README updates",
        ],
    )

    doc.add_heading("Week 2 — defense prep", level=2)
    bullets(
        doc,
        [
            "Run DEFENSE.md demo script twice end-to-end",
            "Prepare answers: class imbalance, false positives, GDPR/cardholder privacy, hybrid vs ML-only",
            "Add 5–10 automated tests",
            "Capture screenshots for written report",
        ],
    )

    doc.add_heading("9. Role capability matrix (for examiners)", level=1)
    table_rows(
        doc,
        ["Feature", "Cardholder", "Analyst", "Admin"],
        [
            ["Own transactions", "Yes", "Limited", "Yes"],
            ["Fraud notifications", "Yes (own)", "Yes", "Yes"],
            ["Report disputes", "Yes", "—", "—"],
            ["Resolve disputes", "—", "Yes", "Yes"],
            ["Global transactions", "No", "Yes", "Yes"],
            ["Fraud analytics", "No", "Yes", "Yes"],
            ["ML model scores", "No", "Yes", "Yes"],
            ["User management", "No", "No", "Yes"],
            ["Fraud rule toggles", "No", "No", "Yes"],
        ],
    )

    doc.add_heading("10. Demo flow (10–15 minutes)", level=1)
    bullets(
        doc,
        [
            "Landing page — problem statement and architecture",
            "Register / login — show RBAC and approval workflow",
            "Admin — user list, toggle a fraud rule",
            "Analyst — start Capture live stream; show toast/sound alerts",
            "Monitoring — filter flagged transactions; approve or flag",
            "Fraud Lab — synthetic high-risk transaction; discuss fused score",
            "Explainability — open transaction ID; walk through feature drivers",
            "Disputes — resolve a cardholder case",
            "Reports — download CSV/PDF",
            "Cardholder login — scoped personal view only",
        ],
    )

    doc.add_heading("11. Bottom line", level=1)
    doc.add_paragraph(
        "The project is above average as a final-year system. To score highly, focus on: (1) making the ML "
        "defensible with a trained artifact and honest metrics, (2) closing end-to-end workflows like disputes, "
        "(3) ensuring explainability works in the live demo, and (4) documenting limitations clearly. "
        "More features without fixing ML credibility can hurt the defense."
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
